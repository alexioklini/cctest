#!/usr/bin/env python3
"""Document-render eval — does write_document turn markdown into the right
structure in .docx and .pdf?

Two assertion classes:
  • BASELINE  — features that already work today (KPI boxes, risk badges, cover,
    TOC, headings without ** / emoji, tables). These MUST stay green across the
    markdown-it refactor — that's how we prove no regression.
  • NEW       — features the current line-parser drops (real bullet/ordered
    lists incl. nested, blockquotes, code blocks, links). These are EXPECTED to
    FAIL on the current converter and turn green after the markdown-it rewrite.

Run:  python3 eval/doc_render_eval.py            # both formats
      python3 eval/doc_render_eval.py --docx     # docx only (no LibreOffice)

The check operates on the produced files directly:
  • docx → python-docx (paragraph styles, numbering, tables, shading XML)
  • pdf  → pdftotext-equivalent via PyMuPDF text + drawing inspection
so it's render-engine-agnostic (doesn't depend on LibreOffice for docx checks).
"""
import sys, os, types, json, re, argparse, zipfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# ── load file_tools with a stubbed brain (no server needed) ────────────────
_brain = types.ModuleType("brain")
_brain.AGENTS_DIR = "/tmp/_doceval_agents"
_brain._after_file_write = lambda *a, **k: None
class _Agent: agent_id = "main"
_brain._current_agent = _Agent()
sys.modules["brain"] = _brain
import engine.context as _ctx
import engine.tools.file_tools as ft
ft._enforce_artifact_path = lambda p, who: (p, None)
ft._brain = _brain

OUT = "/tmp/doc-render-eval"
os.makedirs(OUT, exist_ok=True)

# ── markdown patterns ──────────────────────────────────────────────────────
# Each case: name, markdown, and a list of (assert_kind, format, fn) checks.
# assert_kind ∈ {"baseline","new"}. fn(ctx) -> bool, ctx has docx/pdf probes.

CASES = []
def case(name, md, checks): CASES.append((name, md, checks))

# 1 — headings: no ** / emoji leak (BASELINE)
case("headings_clean",
"""# Hauptbericht
Stichtag: 2025
Verantwortlich: Test

## **📊 Management Summary**
Etwas Text.

## Methodik
Mehr Text.

## Ergebnisse
Noch mehr.

## Anhang
Schluss.
""",
 [("baseline", "docx", lambda c: not c.docx_has_text("**") and "📊" not in c.docx_all_text()),
  ("baseline", "docx", lambda c: c.docx_heading_count() >= 4)])

# 2 — table with risk badges (BASELINE)
case("table_badges",
"""# Risiko

## Faktoren
| Risikofaktor | Gewichtung | Bewertung | Begründung |
|---|---|---|---|
| Cyber | Hoch | Erhöht | x |
| Daten | Mittel | Mittel | y |
| Recht | Niedrig | Gering | z |

## A
## B
## C
""",
 [("baseline", "docx", lambda c: c.docx_table_count() >= 1),
  ("baseline", "docx", lambda c: c.docx_has_shaded_cell()),  # badges shade cells
  ("baseline", "docx", lambda c: not c.docx_has_text("**Risikofaktor**"))])

# 3 — KPI boxes (BASELINE)
case("kpi_boxes",
"""# KPI Bericht
Stichtag: 2025

::kpi 1,55 | Inhärentes Risiko | mittel
::kpi 1,12 | Kontrollumfeld | sehr gut
::kpi 1,34 | Residualrisiko | gering

## Summary
text
## Methodik
text
## Ergebnis
text
""",
 [("baseline", "docx", lambda c: c.docx_has_text("1,55") and c.docx_has_text("Inhärentes Risiko".upper())),
  ("baseline", "docx", lambda c: not c.docx_has_text("::kpi"))])

# 4 — cover + TOC for substantial doc (BASELINE)
case("cover_toc",
"""# Großer Bericht
Stichtag: 2025
Verantwortlich: X

## Eins
text
## Zwei
text
## Drei
text
## Vier
text
""",
 [("baseline", "docx", lambda c: c.docx_has_text("VERTRAULICH") or c.docx_has_text("Inhaltsverzeichnis"))])

# 5 — horizontal rule (BASELINE)
case("hrule",
"""# Doc
Para eins.

---

Para zwei.
""",
 [("baseline", "docx", lambda c: not c.docx_has_text("---"))])

# 6 — bullet list (NEW)
case("bullet_list",
"""# Liste

Vor der Liste.

- Erster Punkt
- Zweiter Punkt
- Dritter Punkt

Nach der Liste.
""",
 [("new", "docx", lambda c: c.docx_has_list()),
  ("new", "docx", lambda c: not c.docx_para_starts_with("- "))])

# 7 — ordered list (NEW)
case("ordered_list",
"""# Nummeriert

1. Eins
2. Zwei
3. Drei
""",
 [("new", "docx", lambda c: c.docx_has_numbered_list()),
  ("new", "docx", lambda c: not c.docx_para_starts_with("1. "))])

# 8 — nested list (NEW)
case("nested_list",
"""# Verschachtelt

- Oben A
  - Unter A1
  - Unter A2
- Oben B
""",
 [("new", "docx", lambda c: c.docx_has_list()),
  ("new", "docx", lambda c: c.docx_list_levels() >= 2)])

# 9 — blockquote (NEW)
case("blockquote",
"""# Zitat

> Dies ist ein Zitat.

Normaler Text.
""",
 [("new", "docx", lambda c: not c.docx_para_starts_with("> ")),
  ("new", "docx", lambda c: c.docx_has_quote_style())])

# 10 — code block (NEW)
case("code_block",
"""# Code

```python
x = 1
print(x)
```

Text danach.
""",
 [("new", "docx", lambda c: not c.docx_has_text("```")),
  ("new", "docx", lambda c: c.docx_has_mono())])

# 11 — link (NEW)
case("link",
"""# Link

Siehe [die Quelle](https://example.com/x) für Details.
""",
 [("new", "docx", lambda c: not c.docx_has_text("](https")),
  ("new", "docx", lambda c: c.docx_has_hyperlink())])

# 12 — mixed real-world (BASELINE structure + NEW lists)
case("mixed",
"""# Komplettbericht
Stichtag: 2025
Verantwortlich: Y

::kpi 2,1 | Risiko | mittel

## Zusammenfassung
Ein Absatz mit **fett** und *kursiv*.

Wichtige Punkte:

- Punkt eins
- Punkt zwei

## Bewertung
| Faktor | Bewertung |
|---|---|
| A | Hoch |
| B | Gering |

## Methodik
text
## Anhang
text
""",
 [("baseline", "docx", lambda c: c.docx_table_count() >= 1 and c.docx_has_shaded_cell()),
  ("baseline", "docx", lambda c: c.docx_has_text("2,1")),
  ("new", "docx", lambda c: c.docx_has_list())])


# ── docx probe ──────────────────────────────────────────────────────────────
class DocxProbe:
    def __init__(self, path):
        import docx
        self.d = docx.Document(path)
        self.z = zipfile.ZipFile(path)
        self._xml = self.z.read("word/document.xml").decode("utf-8", "ignore")
    def docx_all_text(self):
        return "\n".join(p.text for p in self.d.paragraphs) + "\n" + \
               "\n".join(cell.text for t in self.d.tables for row in t.rows for cell in row.cells)
    def docx_has_text(self, s): return s in self.docx_all_text()
    def docx_heading_count(self):
        return sum(1 for p in self.d.paragraphs if p.style.name.startswith("Heading"))
    def docx_table_count(self): return len(self.d.tables)
    def docx_has_shaded_cell(self): return "<w:shd " in self._xml or "<w:shd>" in self._xml
    def docx_para_starts_with(self, pfx):
        return any(p.text.strip().startswith(pfx) for p in self.d.paragraphs)
    def docx_has_list(self):
        # a real list = paragraphs with a numbering reference (numPr) OR List style
        return "<w:numPr>" in self._xml or any("List" in p.style.name for p in self.d.paragraphs)
    def docx_has_numbered_list(self):
        return "<w:numPr>" in self._xml or any("Number" in p.style.name for p in self.d.paragraphs)
    def docx_list_levels(self):
        # Nesting shows up either as numPr ilvl (manual numbering) OR as distinct
        # built-in list styles ('List Bullet' vs 'List Bullet 2') — both are real
        # indent levels in Word, so count whichever the renderer used.
        ilvl = set(re.findall(r'<w:ilvl w:val="(\d+)"', self._xml))
        list_styles = set(p.style.name for p in self.d.paragraphs
                          if "List" in p.style.name)
        return max(len(ilvl), len(list_styles))
    def docx_has_quote_style(self):
        return any("Quote" in p.style.name for p in self.d.paragraphs) or "<w:pBdr>" in self._xml
    def docx_has_mono(self):
        return ("Consolas" in self._xml or "Courier" in self._xml or "<w:shd " in self._xml)
    def docx_has_hyperlink(self):
        return "<w:hyperlink" in self._xml or 'HYPERLINK' in self._xml


class PdfProbe:
    def __init__(self, path):
        import fitz
        self.doc = fitz.open(path)
        self.txt = "\n".join(p.get_text() for p in self.doc)
    def pdf_has_text(self, s): return s in self.txt


def run(do_pdf):
    import importlib
    importlib.reload(ft)  # pick up edits between runs
    ft._enforce_artifact_path = lambda p, who: (p, None)
    ft._brain = _brain
    results = {"baseline": {"pass": 0, "fail": 0, "fails": []},
               "new": {"pass": 0, "fail": 0, "fails": []}}
    with _ctx.request_context():
        g = _ctx.get_request_context(); g.current_agent = _Agent(); g.project = ""
        for name, md, checks in CASES:
            docx_path = f"{OUT}/{name}.docx"
            ft.tool_write_document({"path": docx_path, "content": md, "style": ""})
            probes = {"docx": DocxProbe(docx_path)}
            if do_pdf:
                pdf_path = f"{OUT}/{name}.pdf"
                ft.tool_write_document({"path": pdf_path, "content": md, "style": ""})
                probes["pdf"] = PdfProbe(pdf_path)
            for kind, fmt, fn in checks:
                if fmt == "pdf" and not do_pdf:
                    continue
                ctx = probes[fmt]
                try:
                    ok = bool(fn(ctx))
                except Exception as e:
                    ok = False
                if ok:
                    results[kind]["pass"] += 1
                else:
                    results[kind]["fail"] += 1
                    results[kind]["fails"].append(f"{name}/{fmt}")
    return results


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", action="store_true", help="docx only (skip pdf)")
    args = ap.parse_args()
    r = run(do_pdf=not args.docx)
    print("\n=== BASELINE (must stay green across refactor) ===")
    print(f"  pass {r['baseline']['pass']}  fail {r['baseline']['fail']}")
    if r["baseline"]["fails"]:
        print("  FAILS:", ", ".join(r["baseline"]["fails"]))
    print("=== NEW (fail now, green after markdown-it rewrite) ===")
    print(f"  pass {r['new']['pass']}  fail {r['new']['fail']}")
    if r["new"]["fails"]:
        print("  not-yet:", ", ".join(r["new"]["fails"]))
    # exit code: nonzero only if a BASELINE check fails (that's a regression)
    sys.exit(1 if r["baseline"]["fail"] else 0)
