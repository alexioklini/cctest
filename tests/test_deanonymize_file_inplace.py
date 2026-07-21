"""In-place OOXML deanonymise must not destroy the file (v9.393.1).

Regression for chat 3811cb61: the GDPR after-file-write callback calls
`deanonymize_file(path, path)` — src == dst. The office walker
(`engine.file_pseudonymize._walk_office`) re-zipped by opening `dst_path` in
mode "w", which TRUNCATED the source to 0 bytes before its lazy `zin.read()`
loop had read the members. The read then failed with "Truncated file header",
the re-zip aborted, and a 47 KB .docx was left as a 22-byte empty-EOCD stub —
un-openable in Word — while the model reported a positive content check.

Fix: stage to a sibling temp + atomic os.replace, so the original stays
readable for the whole copy and the swap is all-or-nothing.

Run: python3 -m unittest tests.test_deanonymize_file_inplace -v
"""

from __future__ import annotations

import os
import sys
import tempfile
import unittest
import zipfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pseudonymizer as ps  # noqa: E402


def _needs_docx():
    try:
        import docx  # noqa: F401
        return False
    except ImportError:
        return True


@unittest.skipIf(_needs_docx(), "python-docx not installed")
class TestInPlaceDocxDeanonymise(unittest.TestCase):
    def setUp(self):
        self._maps = []
        self._dir = tempfile.mkdtemp()

    def tearDown(self):
        for m in self._maps:
            ps.close_mapping(m.mapping_id)

    def _mapping_with(self, fake, real):
        m = ps.new_mapping()
        m.forward[real] = fake
        m.reverse[fake] = real
        m.categories[real] = "name"
        self._maps.append(m)
        return m

    def _make_docx(self, path, token, n=50):
        import docx
        doc = docx.Document()
        doc.add_heading("Report", 0)
        for i in range(n):
            doc.add_paragraph(f"Zeile {i}: die Firma {token} wurde geprüft.")
        doc.save(path)
        return os.path.getsize(path)

    def test_inplace_preserves_file_and_reverses(self):
        p = os.path.join(self._dir, "report.docx")
        token = "<NAME_1_abcd>"
        real = "Wiener Privatbank"
        size_before = self._make_docx(p, token)
        m = self._mapping_with(token, real)

        # src == dst — exactly what the after-file-write callback does.
        restored = ps.deanonymize_file(p, p, mapping=m)

        size_after = os.path.getsize(p)
        self.assertGreater(size_after, 1000,
                           f"file was destroyed: {size_after} bytes "
                           f"(was {size_before})")
        # Still a valid zip with the real OOXML members.
        with zipfile.ZipFile(p) as zf:
            names = zf.namelist()
        self.assertIn("word/document.xml", names)
        self.assertGreater(len(names), 5)
        # Every token reversed to the real value.
        self.assertEqual(restored, 50)
        import docx
        text = " ".join(par.text for par in docx.Document(p).paragraphs)
        self.assertIn(real, text)
        self.assertNotIn(token, text)

    def test_no_temp_file_left_behind(self):
        p = os.path.join(self._dir, "clean.docx")
        m = self._mapping_with("<NAME_1_abcd>", "Wiener Privatbank")
        self._make_docx(p, "<NAME_1_abcd>", n=3)
        ps.deanonymize_file(p, p, mapping=m)
        self.assertFalse(os.path.exists(p + ".pii-tmp"),
                         "staging temp file leaked")

    def test_original_survives_when_reverse_raises(self):
        # A corrupt source must NOT be turned into a 22-byte stub — the
        # staging + atomic-replace means the original bytes are untouched
        # on failure (the walker raises, os.replace never runs).
        p = os.path.join(self._dir, "broken.docx")
        with open(p, "wb") as f:
            f.write(b"PK\x03\x04 not really a zip " * 4)
        size_before = os.path.getsize(p)
        m = self._mapping_with("<NAME_1_abcd>", "Wiener Privatbank")
        from engine.file_pseudonymize import FilePseudonymizeError
        with self.assertRaises(FilePseudonymizeError):
            ps.deanonymize_file(p, p, mapping=m)
        self.assertEqual(os.path.getsize(p), size_before,
                         "corrupt source was further mangled")
        self.assertFalse(os.path.exists(p + ".pii-tmp"))


if __name__ == "__main__":
    unittest.main(verbosity=2)
