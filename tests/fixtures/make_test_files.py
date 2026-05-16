"""Generate two synthetic-PII test attachments for transparent-anonymisation
manual testing. Run with the brain-agent venv (python-docx + openpyxl).

Outputs (alongside this script in tests/fixtures/):
  kundenvertrag.docx
  mitarbeiterliste.xlsx

PII is realistic-shape but synthetic — IBANs pass MOD-97, credit cards pass
Luhn, emails / phones / addresses fabricated. Safe to put in any audit
trail.
"""
import os
import docx
from docx.shared import Pt
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ---------------------------------------------------------------------------
# 1) DOCX — Kundenvertrag with embedded PII
# ---------------------------------------------------------------------------
doc = docx.Document()

heading = doc.add_heading("Kundenvertrag Nr. 2026/05/8847", level=1)

doc.add_heading("Vertragsparteien", level=2)
doc.add_paragraph(
    "Auftraggeber: Maria Schneider, geboren am 14.03.1978, "
    "wohnhaft Lindenallee 27, 10437 Berlin, Deutschland. "
    "Kontakt: maria.schneider@beispielmail.de, +49 30 90288471. "
    "Steuer-ID: 81 872 495 633."
)
doc.add_paragraph(
    "Auftragnehmer: Thomas Becker, Geschäftsführer der "
    "Becker Beratung GmbH, Unter den Linden 12, 10117 Berlin. "
    "E-Mail: t.becker@beckerberatung.de, Mobil: +49 171 4488235. "
    "USt-IdNr.: DE298756432."
)

doc.add_heading("Zahlungsdaten", level=2)
doc.add_paragraph(
    "Die monatliche Pauschale in Höhe von 1.450,00 EUR wird "
    "per SEPA-Lastschrift vom Konto des Auftraggebers eingezogen."
)

# Use a real-shape (MOD-97-valid) IBAN: DE89370400440532013000 is a well-known
# test IBAN; valid checksum, no real account.
table = doc.add_table(rows=4, cols=2)
table.style = "Light Grid"
cells = table.rows[0].cells
cells[0].text = "Kontoinhaber"; cells[1].text = "Maria Schneider"
cells = table.rows[1].cells
cells[0].text = "IBAN"; cells[1].text = "DE89 3704 0044 0532 0130 00"
cells = table.rows[2].cells
cells[0].text = "BIC"; cells[1].text = "COBADEFFXXX"
cells = table.rows[3].cells
cells[0].text = "Kreditkarte (Reserve)"; cells[1].text = "4111 1111 1111 1111"

doc.add_heading("Notfallkontakt", level=2)
doc.add_paragraph(
    "Im Falle einer Erkrankung kann der Hausarzt Dr. Klaus Wagner "
    "unter +49 30 33445566 erreicht werden. "
    "Versichertennummer (gesetzliche KV): A123456789."
)

doc.add_heading("Unterschriften", level=2)
doc.add_paragraph(
    "Berlin, den 16.05.2026 — Maria Schneider / Thomas Becker"
)

doc.save(f"{OUT_DIR}/kundenvertrag.docx")
print(f"wrote {OUT_DIR}/kundenvertrag.docx")

# ---------------------------------------------------------------------------
# 2) XLSX — Mitarbeiterliste with embedded PII
# ---------------------------------------------------------------------------
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Mitarbeiter"

headers = ["Name", "Geburtstag", "E-Mail", "Telefon",
           "IBAN (Gehaltskonto)", "Steuer-ID", "Adresse"]
ws.append(headers)
header_font = Font(bold=True, color="FFFFFF")
header_fill = PatternFill("solid", fgColor="2563EB")
for cell in ws[1]:
    cell.font = header_font
    cell.fill = header_fill
    cell.alignment = Alignment(horizontal="left", vertical="center")

rows = [
    ("Anna Becker",        "1985-07-12", "anna.becker@firma.de",
     "+49 30 88445566", "DE89 3704 0044 0532 0130 00",
     "47 123 456 789",  "Schillerstr. 18, 10625 Berlin"),
    ("Markus Hoffmann",    "1979-11-03", "m.hoffmann@firma.de",
     "+49 171 5566778",  "GB82 WEST 1234 5698 7654 32",
     "12 345 678 904",  "Goethestr. 42, 60313 Frankfurt am Main"),
    ("Sofia Lehmann",      "1992-02-25", "s.lehmann@firma.de",
     "+49 89 11223344", "FR76 3000 6000 0112 3456 7890 189",
     "65 901 234 567",  "Maximilianstr. 7, 80539 München"),
    ("Tarek Yılmaz",       "1988-09-30", "t.yilmaz@firma.de",
     "+49 40 99887766", "DE12 5001 0517 0648 4898 90",
     "23 456 789 012",  "Reeperbahn 99, 20359 Hamburg"),
    ("Charlotte Vogel",    "1995-04-18", "c.vogel@firma.de",
     "+49 221 7788995", "AT61 1904 3002 3457 3201",
     "78 901 234 568",  "Hohenzollernring 14, 50672 Köln"),
]
# Cross-row repetition (Anna Becker appears twice) to verify same original
# → same token within a mapping.
rows.append(
    ("Anna Becker",        "1985-07-12", "anna.becker@firma.de",
     "+49 30 88445566", "DE89 3704 0044 0532 0130 00",
     "47 123 456 789",  "Schillerstr. 18, 10625 Berlin"))

for r in rows:
    ws.append(r)

# Width adjustments so the file actually opens readably.
for col_idx, width in enumerate([20, 12, 28, 18, 36, 16, 38], start=1):
    ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = width

# Formula sheet — verify formulas survive (anonymisation skips `=` cells).
ws2 = wb.create_sheet("Gehalt")
ws2.append(["Name", "Brutto", "Steuer", "Netto (Formel)"])
ws2.append(["Anna Becker",      4500, 1350, "=B2-C2"])
ws2.append(["Markus Hoffmann",  5200, 1612, "=B3-C3"])
ws2.append(["Sofia Lehmann",    3800, 1102, "=B4-C4"])

wb.save(f"{OUT_DIR}/mitarbeiterliste.xlsx")
print(f"wrote {OUT_DIR}/mitarbeiterliste.xlsx")
